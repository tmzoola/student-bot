"""Material fayllardan matn ajratish.

PDF (pymupdf/fitz), DOCX (python-docx), TXT (UTF-8) qo'llab-quvvatlanadi.
Rasm-based (skanlangan) PDF uchun OCR yo'q — bo'sh matn qaytaradi.
"""
from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

MIME_PDF = "application/pdf"
MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MIME_TXT = "text/plain"

ALLOWED_MIMES = {MIME_PDF, MIME_DOCX, MIME_TXT}

# 20 MB
MAX_MATERIAL_SIZE = 20 * 1024 * 1024


class ExtractError(Exception):
    """Fayl turini o'qib bo'lmadi yoki bo'sh."""


def _extract_pdf(path: Path) -> str:
    import fitz  # pymupdf

    parts: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    return "\n".join(parts).strip()


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()


def _extract_txt(path: Path) -> str:
    for enc in ("utf-8", "utf-16", "cp1251"):
        try:
            return path.read_text(encoding=enc).strip()
        except UnicodeDecodeError:
            continue
    raise ExtractError("TXT fayl kodlashtirishini aniqlab bo'lmadi")


def extract_text(path: str | Path, mime: str) -> str:
    p = Path(path)
    if not p.exists():
        raise ExtractError(f"Fayl topilmadi: {p}")
    if mime == MIME_PDF:
        return _extract_pdf(p)
    if mime == MIME_DOCX:
        return _extract_docx(p)
    if mime == MIME_TXT:
        return _extract_txt(p)
    raise ExtractError(f"Qo'llab-quvvatlanmaydigan MIME: {mime}")
